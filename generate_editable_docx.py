from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_section(doc, title, content_list):
    doc.add_heading(title, 1)
    for item in content_list:
        p = doc.add_paragraph(item, style='List Bullet')
        run = p.runs[0]
        run.font.name = 'Outfit'
        run.font.size = Pt(12)

def generate_editable_docx():
    doc = Document()
    
    # Title
    title = doc.add_heading('AURAGUIDE PITCH DECK CONTENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Slide 1
    doc.add_heading('Slide 1: Cover', 2)
    doc.add_paragraph('A Lifeline for Family Caregivers.').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Slide 2
    add_section(doc, 'Slide 2: The Caregiver Collapse Cycle', [
        "63M Unpaid Caregivers in the US.",
        "$1T Annual Economic Value of unpaid labor.",
        "87% of caregivers report significant stress.",
        "78% manage complex tasks with zero clinical background."
    ])
    
    # Slide 3
    add_section(doc, 'Slide 3: The Solution — AuraGuide', [
        "Reduce Cognitive Load: Voice-first clinical intelligence.",
        "Navigate Urgency: Real-time crisis response protocols.",
        "Protect Caregiver: Passive burnout detection.",
        "Bridge Language Gaps: 40+ language support."
    ])
    
    # Slide 4
    add_section(doc, 'Slide 4: Market Opportunity', [
        "TAM: $72B US Family Caregiver Support Market.",
        "SAM: $8.5B Digital Caregiver Tech.",
        "10,000 boomers turn 65 daily.",
        "No dominant player in the clinical voice AI space for families."
    ])
    
    # Slide 5
    add_section(doc, 'Slide 5: The Ask — $500,000 Pre-Seed', [
        "40% Product Hardening (App Store Launch)",
        "30% GTM & Beta Ops",
        "30% Team & HIPAA Compliance",
        "Terms: Post-Money SAFE | $3M–$5M Cap"
    ])
    
    output_path = "/Users/vikashvardhan/IdeaProjects/viktronHealthInc/FINAL_Pitchdeck/AuraGuide_Content_Editable_v4.docx"
    doc.save(output_path)
    print(f"Editable DOCX saved to {output_path}")

if __name__ == "__main__":
    generate_editable_docx()
